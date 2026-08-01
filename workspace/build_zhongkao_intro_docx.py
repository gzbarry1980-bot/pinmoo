from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path(r"E:\pinmoo\outputs\zhongkao\广州中考志愿模拟助手-系统详细介绍.docx")

NAVY = "123B66"
BLUE = "2E74B5"
SKY = "EAF3FA"
PALE = "F5F8FB"
GOLD = "E7A629"
GOLD_PALE = "FFF6DF"
GREEN = "2D7C62"
GREEN_PALE = "EAF5F0"
RED = "B54747"
RED_PALE = "FCEEEE"
INK = "263849"
MUTED = "66788A"
GRID = "CBD5DF"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                tag.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            row.cells[idx].width = Inches(width / 1440)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color=GRID, size=8):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("品沐咨询 · 仅供参考    ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_num_definition(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abs_id = max(existing_abs or [0]) + 1
    num_id = max(existing_num or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abs_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def make_numbered(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc, text, num_id):
    p = doc.add_paragraph()
    make_numbered(p, num_id)
    p.paragraph_format.space_after = Pt(7)
    p.add_run(text)
    return p


def add_section_kicker(doc, number, title, lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{number}  ·  SYSTEM GUIDE")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GOLD)
    h = doc.add_heading(title, level=1)
    add_bottom_border(h, BLUE, 10)
    if lead:
        p2 = doc.add_paragraph(lead)
        p2.paragraph_format.space_after = Pt(10)
        p2.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_callout(doc, title, body, fill=SKY, accent=BLUE):
    p = doc.add_paragraph()
    set_paragraph_shading(p, fill)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.16)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(f"{title}\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    r2 = p.add_run(body)
    r2.font.color.rgb = RGBColor.from_string(INK)
    return p


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            if ridx % 2 == 1:
                set_cell_shading(cell, PALE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            r.font.size = Pt(9.2)
            r.font.color.rgb = RGBColor.from_string(INK)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ("List Bullet", "List Bullet 2"):
        s = styles[style_name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(10.3)
        s.font.color.rgb = RGBColor.from_string(INK)
        s.paragraph_format.space_after = Pt(5)
        s.paragraph_format.line_spacing = 1.12

    heading_specs = {
        "Title": (30, NAVY, 0, 14),
        "Subtitle": (13, MUTED, 0, 8),
        "Heading 1": (17, NAVY, 13, 7),
        "Heading 2": (13, BLUE, 10, 5),
        "Heading 3": (11.5, NAVY, 7, 3),
    }
    for name, (size, color, before, after) in heading_specs.items():
        s = styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "广州中考志愿模拟助手｜系统详细介绍"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_bottom_border(hp, GRID, 6)

    footer = section.footer
    add_page_number(footer.paragraphs[0])
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("产品介绍  |  2026版")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(GOLD)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(24)
    title.add_run("广州中考\n志愿模拟助手")

    sub = doc.add_paragraph(style="Subtitle")
    sub.add_run("把分数、梯度、志愿顺序与历史数据，\n放进同一套可解释的模拟流程。")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("输入分数，先模拟，再填报")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    add_bottom_border(p, GOLD, 16)

    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(tbl, [3120, 3120, 3120])
    set_repeat_table_header(tbl.rows[0])
    metrics = [
        ("3种", "使用模式"),
        ("15个", "普通高中志愿槽"),
        ("10,000次", "未来情景模拟"),
    ]
    for i, (big, small) in enumerate(metrics):
        set_cell_shading(tbl.cell(0, i), SKY if i != 1 else GOLD_PALE)
        set_cell_shading(tbl.cell(1, i), SKY if i != 1 else GOLD_PALE)
        for c in (tbl.cell(0, i), tbl.cell(1, i)):
            set_cell_margins(c, top=120, bottom=120)
            set_cell_border(c, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
        p1 = tbl.cell(0, i).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p1.add_run(big)
        rr.bold = True
        rr.font.size = Pt(19)
        rr.font.color.rgb = RGBColor.from_string(NAVY)
        p2 = tbl.cell(1, i).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr2 = p2.add_run(small)
        rr2.font.size = Pt(9.5)
        rr2.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(35)
    r = p.add_run("由品沐咨询提供")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = doc.add_paragraph("版本日期：2026年7月26日")
    p2.runs[0].font.size = Pt(9.5)
    p2.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    p3 = doc.add_paragraph()
    add_hyperlink(p3, "zhongkao.pinmooconsulting.com", "https://zhongkao.pinmooconsulting.com/")
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    num_id = add_num_definition(doc)
    add_cover(doc)

    add_section_kicker(
        doc,
        "01",
        "不是替家长做决定，而是把决定的后果先模拟一遍",
        "同样的分数，因为批次、志愿顺序、梯度线、资格与学校热度不同，最终去向可能完全不同。系统的价值，是让家长在正式提交前看见风险、机会与可调整空间。",
    )
    add_callout(
        doc,
        "家长最容易忽略的不是“能不能报”，而是“这样排序之后会发生什么”。",
        "高分并不等于每个志愿都安全；分数达到学校最低线，也不代表一定能以当前志愿序号被录取。系统会把梯度、志愿序号、资格限制和历史波动放在一起判断。",
        GOLD_PALE,
        RED,
    )
    doc.add_heading("产品一句话定位", level=2)
    doc.add_paragraph(
        "“广州中考志愿模拟助手”是一套面向广州中考家庭的志愿决策工具。家长可以从大概分数出发生成方向，也可以围绕目标学校倒推冲刺分值，或把已有志愿表交给系统复核。系统输出方案评分、单校机会区间、最可能录取学校、未被普通高中录取的估算风险，以及可直接执行的调整建议。"
    )
    doc.add_heading("它解决的四个典型问题", level=2)
    for text in [
        "学校很多、规则复杂，不知道从哪里开始选。",
        "心里已有方案，但不确定冲、稳、保是否真正拉开。",
        "有目标学校，却不知道需要冲到多少分、应该放在哪个志愿位置。",
        "只看往年最低分，忽略了末位志愿序号、梯度保护、资格与年度波动。",
    ]:
        add_bullet(doc, text)
    add_table(
        doc,
        ["产品要素", "当前说明"],
        [
            ["适用对象", "2026年广州中考考生家庭；亦可用于后续年度的规则暂估与方向讨论"],
            ["覆盖批次", "普通高中第二批次3个、第三批次6个、第四批次6个志愿"],
            ["核心输出", "方案评分、冲稳保结构、录取机会区间、置信度、最可能录取去向与具体调整建议"],
            ["使用方式", "网页工具；匿名使用；支持本地草稿、JSON导入导出、打印与另存PDF"],
            ["解锁方式", "9.99元序列号；一个序列号最多绑定2台不同设备"],
        ],
        [2100, 7260],
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "02",
        "三种入口，对应三类真实填报状态",
        "系统不要求每位家长一开始就懂学校。先判断家长现在掌握了什么，再进入相应流程。三种模式可以互相切换，最终都能形成一份可继续修改的模拟志愿表。",
    )
    add_table(
        doc,
        ["使用模式", "适合谁", "需要填写", "系统给什么"],
        [
            ["方向推荐", "只有大概估分、尚未研究学校", "估分上下限、考生类别、升学区、户籍区、风险偏好等", "优先参考学校录取门槛位次与户籍所在区，生成进取／均衡／稳健三套冲稳保方向草案"],
            ["目标学校", "已有心仪学校，希望知道怎么冲", "目标学校、当前估分、考生类别、升学区与资格信息", "建议冲刺分值、历史门槛、适合的志愿位置、上下衔接学校与风险提示"],
            ["方案求证", "已经有志愿计划，想检查是否合理", "完整考生条件与自己的15个志愿", "100分评分、单校机会、录取去向分布、风险项和可自动应用的补强建议"],
        ],
        [1500, 2080, 2740, 3040],
    )
    doc.add_heading("方向推荐：不研究学校也能先有一个可讨论的起点", level=2)
    doc.add_paragraph(
        "家长只需输入估分下限和上限，再补充考生类别、升学区域、户籍所在区与希望方案更偏进取、均衡还是稳健。系统先筛除明显不符合报考范围或偏好的学校，再按历史门槛位次、波动程度、同区优先与志愿层级形成草案。"
    )
    add_callout(
        doc,
        "切换风险偏好后，变化的不只是提示文字。",
        "进取方案会保留更多上探学校；均衡方案强调冲、稳、保的层次；稳健方案会提高可接受保底的数量和位置。系统会重新计算机会结构、预估去向与未录取风险。",
        GREEN_PALE,
        GREEN,
    )
    doc.add_heading("目标学校：从“想去”倒推“要做到什么”", level=2)
    doc.add_paragraph(
        "选择目标学校后，系统会结合其历年录取门槛、招生计划变化、同分序号与末位志愿序号，给出目标分值区间与填报位置建议。若当前估分差距较大，会同时提供同层级替代方向和更稳妥的承接学校。"
    )
    doc.add_heading("方案求证：对已有志愿表进行结构化体检", level=2)
    doc.add_paragraph(
        "系统检查重复学校、空档、资格冲突、批次与志愿顺序，并用同一组模拟情景计算所有学校的去向分布。家长可以点击“去调整”，由系统直接应用推荐调整，再重新分析，而不需要逐项手工改表。"
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "03",
        "一次完整使用，从考生条件到可打印报告",
        "系统将复杂规则拆成连续步骤，减少来回跳转。页面顶端与快捷导航均可返回首页、切换模式或进入专项信息页。",
    )
    steps = [
        ("填写考生信息", "选择目标年度、复盘或预测模式，输入分数或估分区间、户籍与学籍区、考生类别及必要资格。"),
        ("设定选校偏好", "选择公民办、学费上限、住宿、区域、排除学校与风险偏好；同机会档优先推荐户籍同区学校。"),
        ("筛选或生成学校", "可以搜索学校，也可以由系统按方向自动生成冲稳保草案。特长生、自主招生等从独立入口查询。"),
        ("完成模拟志愿表", "按2026结构填写第二批3个、第三批6个、第四批6个志愿；支持排序、重复拦截与连续志愿校验。"),
        ("运行方案分析", "系统先做资格与表格检查，再运行历史复盘或未来情景模拟，计算方案评分和各校机会。"),
        ("查看去向与建议", "读取最可能录取学校、各校机会区间、未录取风险和具体批次／志愿序号的调整建议。"),
        ("保存、导出与复核", "草稿保存在当前浏览器，可匿名导出JSON；最终可打印或由浏览器另存为PDF，并在正式填报前对照官方信息。"),
    ]
    for title, body in steps:
        p = doc.add_paragraph()
        make_numbered(p, num_id)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(title + "｜")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(body)

    doc.add_heading("需要填写哪些信息", level=2)
    add_table(
        doc,
        ["类别", "主要字段", "用途"],
        [
            ["成绩", "分数／估分上下限、同分序号", "定位分数梯度、转换历史位次、判断同分情形"],
            ["身份与区域", "户籍区、学籍区、考生类别、是否跨区", "判断招生范围、外区计划与可报资格"],
            ["名额分配", "所在初中、名额分配资格、是否参加第二批", "匹配第二批分配学校、计划与历年录取结果"],
            ["学业条件", "参考科目等级等", "核验学校或类别的报考要求"],
            ["家庭偏好", "公民办、学费、住宿、区域、排除学校、风险偏好", "让建议在录取机会之外也符合家庭可接受范围"],
        ],
        [1560, 3600, 4200],
    )
    add_callout(
        doc,
        "隐私边界",
        "系统不要求填写姓名、准考证号、身份证号或联系方式。草稿默认仅保存在当前浏览器；匿名导出前会再次提示检查是否包含个人身份信息。",
        SKY,
        NAVY,
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "04",
        "广州普通高中录取逻辑，系统如何理解",
        "模拟的基础不是“某校去年多少分”，而是广州中考在批次内按梯度、志愿顺序和成绩择优投档的组合规则。不同年度政策以当年官方文件为准。",
    )
    doc.add_heading("核心原则：梯度投档、志愿优先、择优录取", level=2)
    add_number(doc, "先按批次处理。前一批次一旦录取，后续批次不再继续投档。", num_id)
    add_number(doc, "批次内按投档控制线形成分数梯度；高梯度考生在投向低梯度学校时可能获得梯度保护。", num_id)
    add_number(doc, "同一梯度内重视志愿序号。学校若在第一志愿已完成计划，后续志愿即使达到最低分，也可能没有机会。", num_id)
    add_number(doc, "同一志愿序号内再按成绩择优；同分时需要结合官方公布的同分序号或相关排序规则。", num_id)
    add_callout(
        doc,
        "最低录取分数线不是一张“入场券”。",
        "它是当年实际投档结果的截面。家长还要同时看该校在什么志愿序号完成录取、考生所在梯度、同分排序、招生范围和当年竞争变化。",
        RED_PALE,
        RED,
    )
    doc.add_heading("示意案例：715分为什么可能到第三志愿才录取", level=2)
    add_table(
        doc,
        ["志愿", "学校与示意门槛", "模拟判断"],
        [
            ["第一志愿", "铁一中学越秀校区｜717分", "考生715分未达到示意门槛，不能投档"],
            ["第二志愿", "广州市第六中学海珠校区｜712分", "分数达到，但若当年该校在第一志愿已经完成计划，第二志愿仍无法录取"],
            ["第三志愿", "清华附中湾区学校｜698分", "考生处于更高梯度，学校位于较低梯度；在符合当年投档条件时，可通过梯度保护获得录取机会"],
        ],
        [1400, 3200, 4760],
    )
    p = doc.add_paragraph("说明：以上为规则理解示意，学校名称、分数与志愿结果不能脱离对应年度、批次、招生范围和官方录取表单独使用。")
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_heading("历史复盘的三种判断", level=2)
    add_table(
        doc,
        ["结果", "含义", "需要注意"],
        [
            ["可投档", "按该年度已公开条件，考生在此志愿具备投档可能", "仍应对照招生范围、计划类型及最终官方录取结果"],
            ["不能投档", "分数、梯度、志愿序号或资格中至少一项不满足", "系统会指出具体原因"],
            ["结果不确定", "分数相同但同分序号不足，或官方数据缺少关键字段", "不会强行判定为录取"],
        ],
        [1600, 4100, 3660],
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "05",
        "历史复盘与未来预测，是两种不同的答案",
        "复盘回答“如果当年这样填，可能发生什么”；预测回答“把历史波动带入未来后，各种结果出现的频率如何”。两者不能混为一谈。",
    )
    doc.add_heading("历史复盘模式", level=2)
    doc.add_paragraph(
        "选择2021—2026年后，系统按该年度可用的梯度线、学校录取线、末位志愿序号和同分序号进行复演。若某年度或某批次官方数据缺失，系统应明确显示数据不足，而不是用邻近年度补齐成“真实结果”。"
    )
    doc.add_heading("未来预测模式", level=2)
    doc.add_paragraph(
        "面向2027年及以后，在新年度政策尚未正式发布前，页面明确显示“暂按2026规则模拟”。未来预测不是简单取历年平均分，而是把估分波动、历史年份差异、分数位次和批次投档顺序放进10,000次确定性情景模拟。"
    )
    add_table(
        doc,
        ["环节", "处理方式"],
        [
            ["估分", "按用户填写的下限—中心值—上限形成三角分布，兼顾发挥不理想、正常和较好情形"],
            ["历史权重", "2021—2026年依次采用5%、8%、12%、20%、25%、30%，让近年数据权重更高"],
            ["位次换算", "通过当年分数段将分数转换为可比较的位次，再映射学校门槛"],
            ["投档顺序", "逐批次、逐志愿模拟“梯度优先、同梯度志愿优先、分数择优”，首个满足条件的学校作为该情景结果"],
            ["机会结果", "输出机会中值与区间；结果统一限制在5%—95%，不出现0%或100%的保证式表达"],
            ["置信度", "同口径数据5年以上且波动较小为高；3—4年为中；不足3年或计划类型变化为低"],
        ],
        [1960, 7400],
    )
    add_callout(
        doc,
        "为什么不把各校概率简单相加？",
        "同一个考生只能在同一情景中得到一个最终去向。系统用同一组模拟情景统计“最可能录取学校”“各校结果占比”和“未被普通高中录取风险”，避免多个单校概率相加后超过100%。",
        GREEN_PALE,
        GREEN,
    )
    doc.add_heading("单校机会档位", level=2)
    add_table(
        doc,
        ["档位", "机会中值", "正确理解"],
        [
            ["冲刺", "低于45%", "有上探空间，但结果对分数、志愿热度和梯度变化更敏感"],
            ["匹配", "45%—74%", "与当前条件较接近，仍需用后续志愿承接风险"],
            ["保底", "75%及以上", "相对稳妥，但不是录取承诺，仍受资格和当年竞争影响"],
        ],
        [1500, 1800, 6060],
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "06",
        "100分方案评分：告诉家长哪里需要改",
        "评分的目的不是制造“漂亮数字”，而是检查志愿表是否有效、是否充分利用槽位、是否形成层次、顺序是否保护分数，以及有没有家庭真正能接受的保底。",
    )
    add_table(
        doc,
        ["评分维度", "分值", "检查重点"],
        [
            ["资格与表格有效性", "20分", "招生范围、考生类别、名额分配、参考科目、重复学校、连续志愿等"],
            ["志愿槽位利用", "15分", "第二、第三、第四批次的可用位置是否被合理利用，是否出现无必要空档"],
            ["冲稳保结构", "25分", "是否真正形成机会层次，避免所有学校都集中在同一风险区间"],
            ["志愿顺序与梯度保护", "25分", "学校排序是否符合意愿，是否兼顾末位志愿序号和梯度保护"],
            ["保底完整性及偏好匹配", "15分", "是否有可接受的保底，学费、住宿、区域、公民办等是否符合家庭限制"],
        ],
        [2460, 1200, 5700],
    )
    doc.add_heading("分数怎么解释", level=2)
    add_table(
        doc,
        ["总分", "评价", "建议动作"],
        [
            ["85—100", "稳健", "整体结构完整，重点复核资格、偏好和最新官方数据"],
            ["70—84", "基本合理", "存在局部次序、槽位或保底问题，建议按系统提示优化"],
            ["55—69", "风险偏高", "冲稳保失衡或承接不足，应调整具体批次与志愿位置"],
            ["低于55", "建议重排", "先解决资格与表格问题，再重新建立冲稳保结构"],
        ],
        [1700, 2000, 5660],
    )
    add_callout(
        doc,
        "评分设有安全上限",
        "出现无资格学校时，总分最高59分；没有家庭可接受的保底，或学校学费明确超过家庭上限时，总分最高69分。系统不会用其他加分项掩盖关键风险。",
        RED_PALE,
        RED,
    )
    doc.add_heading("“满分补强清单”如何工作", level=2)
    doc.add_paragraph(
        "分析后，系统按优先级列出离100分还缺什么，例如补齐可用志愿、增加匹配校、补充可接受保底、调整高低梯度顺序或移除无资格学校。建议必须指出具体批次和志愿序号，并给出可替换学校、调整原因以及调整后的预计变化。点击“去调整”后，系统可直接应用选择并重新分析。"
    )
    add_table(
        doc,
        ["家长看到的结果", "应该怎么读"],
        [
            ["方案评分", "看志愿表结构和规则风险，不等同于考生成绩或录取概率"],
            ["机会区间", "看不确定性范围，不只看中间值；区间越宽，年度波动或数据不足越明显"],
            ["置信度", "判断依据是否充足、口径是否稳定；低置信度结果需要更多人工核实"],
            ["最可能录取学校", "是10,000次情景中出现频率较高的去向，不是录取承诺"],
            ["未录取风险", "提示普通高中志愿整体承接是否足够，不能用单个保底标签替代"],
        ],
        [2900, 6460],
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "07",
        "学校信息不只给分数，还要说明依据与边界",
        "系统尽量把学校、校区、招生类别和不同考生口径拆开，避免把不同年份、不同批次或不同招生范围的数据混成一个数字。",
    )
    doc.add_heading("当前数据概况", level=2)
    add_table(
        doc,
        ["项目", "数量／范围", "说明"],
        [
            ["学校主表", "311所", "按学校与校区建立标准化标识，学校更名或校区拆分需保留映射"],
            ["招生录取记录", "1,728条", "覆盖已解析的年度、批次、招生类别、分数线、同分序号与末位志愿序号"],
            ["第二批名额分配记录", "27,222条", "细化到考生所在初中、分配学校、计划数及历年结果"],
            ["分数段记录", "180条", "用于跨年度位次换算与预测模拟"],
            ["来源清单", "27项", "记录来源URL、抓取时间、文件哈希与解析版本，便于追溯"],
            ["历史年度", "2021—2026", "以官方公开资料为主；不同年份的可解析完整度不同"],
        ],
        [2200, 1700, 5460],
    )
    p = doc.add_paragraph("数据版本：20260722-d38b3ec0｜生成时间：2026-07-22｜本介绍制作时读取的版本快照")
    p.runs[0].font.size = Pt(8.7)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_heading("数据字段", level=2)
    doc.add_paragraph(
        "标准记录包括学校／校区、区域、性质、招生范围、年度、批次、招生类别、户籍／随迁／外区口径、招生计划、实际录取人数、录取分数线、同分序号、末位志愿序号、梯度线、分数段与来源。第二批名额分配进一步记录所在初中和分配学校。"
    )
    doc.add_heading("明确不做的推定", level=2)
    for text in [
        "实际录取人数没有统一官方依据时保留为空，不用计划数代替实际录取数。",
        "学校自报高考成绩、升学率只作为独立参考卡片，标注年份、统计口径和来源，不进入录取模型。",
        "不输出主观“学校质量总榜”；仅提供录取门槛位次、招生规模、波动程度等可核验指标。",
        "住宿、学费、开放日、自主招生等动态信息必须标明来源和更新时间；不能确认时提示家长向学校核实。",
    ]:
        add_bullet(doc, text)
    add_callout(
        doc,
        "数据缺口会降低置信度，而不是被隐藏。",
        "2021年第三、第四批次及2022年第三批次部分官方页面为图片形式，自动化数据集中保留来源并将缺失记录明确标记。任何缺年、学校更名、校区拆分或计划大幅变化，都应触发置信度下调。",
        GOLD_PALE,
        GOLD,
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "08",
        "专项信息独立呈现：特长生、自主招生与名额分配",
        "这些类别的资格、学校要求和沟通节点与普通高中统一招生不同，因此系统提供查询与决策辅助，但不把它们混进第三、第四批次的普通录取概率。",
    )
    doc.add_heading("特长生：先按项目找学校", level=2)
    doc.add_paragraph(
        "家长可以先选择足球、篮球、田径、艺术等特长类别，再查看开设相应项目的学校、往年资格要求、招生计划和参考分值。参考分值必须说明它来自普通批次门槛、专项最低控制要求还是历史公开材料，不能把固定比例换算成看似精确的“学校录取分”。"
    )
    doc.add_heading("自主招生：先选学校，再核验状态", level=2)
    doc.add_paragraph(
        "自主招生资格往往与具体学校、报名材料、综合能力考核和年度安排相关。正确流程是先选择学校，再展示该校当年或往年的招生简章、报名条件、开放日／校园咨询日、官方联系方式与资料准备提示。系统不在未选择学校时提前判断“资格通过”。"
    )
    doc.add_heading("学校沟通信息怎么用", level=2)
    add_table(
        doc,
        ["信息", "用途", "使用边界"],
        [
            ["开放日／校园咨询日", "了解培养方向、招生要求、课程与住宿；提前准备问题清单", "每年安排可能变化，以学校官方通知为准"],
            ["校庆日", "可作为了解学校文化与公开活动的参考节点", "不能推定为自主招生沟通或录取节点"],
            ["学校电话／官网／公众号", "核实报名、材料、资格、考试和时间安排", "优先使用学校官方渠道，避免非官方中介信息"],
            ["参考分值", "帮助判断是否值得重点准备或同时设置普通批次承接", "不得理解为专项录取保证线"],
        ],
        [2100, 3600, 3660],
    )
    add_callout(
        doc,
        "第二批名额分配需要考生主动确认。",
        "只有勾选“我确认参加第二批名额分配”后，相关志愿与字段才显示；取消勾选时同步取消相关选择，避免家长误以为系统默认具备资格。",
        SKY,
        NAVY,
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "09",
        "解锁、设备与隐私：一次购买，长期保留使用权",
        "正式版采用序列号解锁，不在页面内直接展示支付入口。用户从购买平台取得序列号后，在解锁页完成设备绑定。",
    )
    add_table(
        doc,
        ["项目", "规则"],
        [
            ["产品价格", "9.99元／个序列号；实际售卖价格与活动规则以购买平台页面为准"],
            ["设备数量", "同一序列号最多绑定2台不同设备；第3台设备需要新的序列号"],
            ["再次访问", "已成功解锁的设备再次进入时显示“已解锁”、已绑定设备数、设备上限与剩余可绑定次数"],
            ["服务端记录", "正式序列号与设备绑定状态由服务端保存，避免只依赖浏览器一次性状态"],
            ["异常处理", "若无法解锁或序列号状态异常，请联系原购买平台客服并提供订单信息查询"],
        ],
        [2200, 7160],
    )
    doc.add_heading("为什么仍然不需要账号", level=2)
    doc.add_paragraph(
        "序列号用于验证使用权，不等于用户账号。志愿草稿仍以当前浏览器本地保存为主，系统不要求绑定姓名、手机号、微信或准考证。更换设备时，家长可以通过匿名JSON导出与导入迁移自己的草稿。"
    )
    doc.add_heading("建议用户保存的三样东西", level=2)
    add_number(doc, "购买订单记录：出现序列号问题时用于向原购买平台客服查询。", num_id)
    add_number(doc, "序列号：不要公开发布或转发给无关人员，以免消耗设备绑定名额。", num_id)
    add_number(doc, "匿名志愿草稿：在重要调整后导出一份JSON备份；导出前确认没有自行填写身份信息。", num_id)
    add_callout(
        doc,
        "访问地址",
        "备案、域名解析和服务器配置完成后，以 https://zhongkao.pinmooconsulting.com/ 为正式入口。若页面显示维护或备案提示，请以品沐咨询发布的最新通知为准。",
        GREEN_PALE,
        GREEN,
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "10",
        "家长如何正确使用结果",
        "最好的使用方式不是追求一个“满分志愿表”，而是通过多轮模拟看清：分数变化后，哪些学校最敏感；志愿顺序变化后，哪些风险被放大；家庭能否接受真正的保底去向。",
    )
    doc.add_heading("推荐的家庭决策流程", level=2)
    recommendations = [
        "先用方向推荐生成一版均衡草案，建立大致学校范围。",
        "把学校按“最想去、可以接受、只作保底、明确不接受”分组，补充学费、住宿和区域限制。",
        "对1—3所核心目标学校使用目标学校模式，比较所需分值和合适志愿位置。",
        "将自己的最终想法放入方案求证，查看资格、梯度、志愿序号和未录取风险。",
        "分别用发挥不理想、正常和较好三个分数情景复核，不只测试一个中心分。",
        "保存两套方案：一套更贴近真实意愿，一套在成绩公布后可以迅速切换。",
        "正式提交前，逐项对照当年官方报考指南、学校招生简章和中考服务平台。",
    ]
    for text in recommendations:
        add_number(doc, text, num_id)
    doc.add_heading("常见误区", level=2)
    add_table(
        doc,
        ["误区", "正确做法"],
        [
            ["只按去年最低分从高到低排", "同时看位次、梯度、末位志愿序号、计划变化和招生范围"],
            ["分数达到学校线就一定能进", "核对学校在哪个志愿序号完成录取，以及同分序号是否满足"],
            ["所有志愿都填名气更大的学校", "保留真实可接受的匹配与保底，避免机会结构全部集中在冲刺档"],
            ["评分越高越代表学校越好", "评分只评价志愿方案结构，不评价学校质量，也不替代家庭偏好"],
            ["概率80%就是基本保证", "任何概率都是统计估计；规则、计划、成绩分布和竞争热度都可能变化"],
            ["专项报名后普通批次不用准备", "专项资格与录取存在不确定性，应同时准备普通高中志愿承接"],
        ],
        [3600, 5760],
    )

    doc.add_page_break()
    add_section_kicker(
        doc,
        "11",
        "常见问题",
        "以下回答用于帮助家长理解产品边界。涉及个人资格和当年政策时，仍需向初中学校或招考部门核实。",
    )
    faqs = [
        ("系统会直接告诉我“肯定能上哪所学校”吗？", "不会。系统提供机会区间、档位与置信度，不显示100%录取，也不作任何录取承诺。"),
        ("2026年的录取结果是否已经全部收录？", "数据集会按官方公布进度更新。只有官方已经发布并通过校验的字段才能进入正式版本；尚未公布或无法可靠解析的内容会显示缺失或低置信度。"),
        ("为什么同一所学校不同考生看到的建议不一样？", "考生类别、户籍与学籍区、跨区、名额分配资格、志愿位置、估分区间和家庭偏好都会改变可报范围与模拟结果。"),
        ("为什么某校分数比我低，系统仍提示风险？", "学校可能在更早志愿已经完成计划，或考生受到招生范围、梯度、同分序号和计划类型限制。最低分不能单独决定结果。"),
        ("学校排名按什么排？", "系统优先使用可核验的录取门槛位次、招生规模和波动程度，不提供主观学校质量总榜。方向推荐在同机会档内优先兼顾户籍所在区。"),
        ("高考出口率、重本率会进入推荐吗？", "不会直接进入录取模型。学校自报的高考成绩或升学率口径不统一，只能作为注明年份和来源的独立参考，不能直接横向比较。"),
        ("换手机或电脑后还能用吗？", "可以。同一序列号最多绑定2台不同设备；达到上限后，第3台设备需要新的序列号。志愿草稿建议通过匿名JSON自行迁移。"),
        ("系统能代替学校老师或官方填报平台吗？", "不能。它用于信息整理、情景模拟和风险提示，正式填报必须在官方平台完成，并以当年政策和学校通知为准。"),
    ]
    doc.add_heading("问题与回答", level=2)
    for q, a in faqs:
        h = doc.add_heading(q, level=3)
        h.paragraph_format.keep_with_next = True
        p = doc.add_paragraph(a)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()
    add_section_kicker(
        doc,
        "12",
        "官方依据与免责声明",
        "系统以广州市招生考试委员会办公室、广州市教育局及中考服务平台的正式发布为最高依据。以下链接便于家长核对规则与原始数据。",
    )
    sources = [
        ("2026年广州市高中阶段学校招生报考指南及志愿模拟表", "https://gzzk.gz.gov.cn/attachment/8/8027/8027351/10816553.pdf"),
        ("2026年广州市中考志愿填报问答", "https://gzzk.gz.gov.cn/gkmlpt/content/10/10811/post_10811628.html?jump=true"),
        ("广州市中考历年录取分数栏目", "https://gzzk.gz.gov.cn/zkzz/zkxx/lnfs/index.html"),
        ("2026年广州市高中阶段学校招生录取控制线", "https://gzzk.gz.gov.cn/zkzz/zkxx/wjtz/content/post_10907243.html"),
    ]
    for title, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, title, url)
        p.add_run(f"\n{url}").font.size = Pt(8.3)
        p.runs[-1].font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_heading("完整免责声明", level=2)
    disclaimer = (
        "本系统依据公开招生政策及历史数据进行模拟分析，所示评分、录取机会和学校建议均为统计估计，不代表官方录取结果或任何录取承诺。招生政策、计划、报考范围、成绩分布和志愿竞争每年可能变化，请以当年广州市教育局、广州市招生考试委员会办公室及中考服务平台最终公布的信息为准。名额分配、随迁子女、跨区及其他资格请向学校或招考部门核实。志愿选择由考生及监护人自行决定。本系统仅供参考。"
    )
    add_callout(doc, "请在正式填报前再次阅读", disclaimer, GOLD_PALE, RED)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("广州中考志愿模拟助手")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = doc.add_paragraph("由品沐咨询提供｜让每一次志愿调整都有依据")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    props = doc.core_properties
    props.title = "广州中考志愿模拟助手｜系统详细介绍"
    props.subject = "产品功能、录取逻辑、数据口径、使用流程与免责声明"
    props.author = "品沐咨询"
    props.keywords = "广州中考, 志愿填报, 志愿模拟, 梯度保护, 品沐咨询"
    props.comments = "2026版"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
